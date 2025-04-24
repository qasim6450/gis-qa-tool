import streamlit as st
import os
import csv
import pandas as pd
import geopandas as gpd
from datetime import datetime
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import re
from difflib import SequenceMatcher
from typing import List, Dict, Set, Tuple, Optional, Union
from spellchecker import SpellChecker
import base64
from pathlib import Path
import tempfile
import shutil
import zipfile

spell = SpellChecker()

# Set page config with icon
st.set_page_config(
    page_title="Zoning Data QA Tool",
    page_icon="zoneomics_icon.png",  # or "🏙️" as fallback
    layout="wide"
)

def cleanup_temp_dir(temp_dir):
    """Remove temporary directory and its contents"""
    if temp_dir and os.path.exists(temp_dir):
        shutil.rmtree(temp_dir)

if 'temp_dir' in st.session_state and st.session_state.temp_dir:
    # Clean up previous temp directory if exists
    cleanup_temp_dir(st.session_state.temp_dir)
    del st.session_state.temp_dir

# Custom CSS
# Add this CSS at the top of your script (with other CSS)
st.markdown("""
<style>
    .header-container {
        display: flex;
        align-items: center;
        gap: 0.1px;
        margin: 0 0 1.5rem 0;
        padding: 0;
    }
    .header-logo {
        height: 80px;  /* Fixed height for perfect vertical alignment */
        width: auto;
        object-fit: contain;
        margin-top: -4px; /* Fine-tune vertical position */
    }
    .header-title {
        margin: 0;
        padding: 0;
        line-height: 1.2;
    }
</style>
""", unsafe_allow_html=True)
st.markdown("""
<style>
    .header-style { font-size:24px; font-weight:bold; color:#2E86C1; margin:20px 0 10px; }
    .subheader-style { font-size:18px; font-weight:bold; color:#3498DB; margin-top:15px; }
    .error-style { color:#E74C3C; font-weight:bold; }
    .success-style { color:#27AE60; font-weight:bold; }
    .warning-style { color:#F39C12; font-weight:bold; }
    .info-box { background-color:#EBF5FB; padding:10px; border-radius:5px; margin-bottom:10px; }
    .zone-code-display { font-family:monospace; white-space:pre-wrap; background-color:#f5f5f5; padding:10px; border-radius:5px; }
    .duplicate-header { color:#E74C3C; font-weight:bold; font-size:20px; margin-top:20px; }
    .step-status { font-weight:bold; margin-left:10px; }
    .footer {
        position: fixed;
        right: 0;
        bottom: 0;
        width: auto;
        padding: 10px;
        color: gray;
        font-size: 1.5em;}
</style>
""", unsafe_allow_html=True)

# Initialize session state
if 'qa_results' not in st.session_state:
    st.session_state.qa_results = {}
if 'current_folder' not in st.session_state:
    st.session_state.current_folder = None
if 'plu_files_corrected' not in st.session_state:
    st.session_state.plu_files_corrected = False

VALID_ZONES = {
    'Residential': ['Single Family', 'Two Family', 'Multi Family'],
    'Agriculture': ['Agriculture'],
    'Commercial': ['General Commercial', 'Retail Commercial', 'Special Commercial', 
                   'Neighborhood Commercial', 'Core Commercial', 'Office'],
    'Industrial': ['Light Industrial', 'Industrial'],
    'Special': ['Special'],
    'Overlay': ['Overlay'],
    'Planned': ['Planned'],
    'Mixed': ['Mixed Use']
}

FULL_REQUIRED_COLUMNS = [
    'zone_code', 'zone_name', 'zone_type', 'zone_sub_type',
    'zone_code_link', 'structure_type', 'automatic_scraping',
    'source', 'format', 'ordinance_publishing_date',
    'map_publishing_date', 'zone_guide'
]

ZONEOMICS_REQUIRED_COLUMNS = [
    'zone_code', 'zone_name', 'zone_type', 'zone_sub_type',
    'zone_code_link', 'ordinance_publishing_date',
    'map_publishing_date', 'zone_guide'
]

STANDARD_CONTROLS_COLUMNS = [
    "zone_code", "min_lot_area_sq_ft", "min_lot_width_ft",
    "max_building_height_ft", "max_far", "max_coverage",
    "max_impervious_coverage_percentage", "min_front_yard_ft",
    "min_side_yard_ft", "min_side_yard_at_least_one_ft",
    "min_side_yard_both_ft", "min_rear_yard_ft",
    "min_landscaped_space_percentage", "open_space_percentage",
    "private_open_space_percentage", "max_density_du_per_acre"
]

PLU_USE_TYPES = {
    'S': 'Special Uses:- S',
    'P': 'Permitted Uses:- P',
    'A': 'Accessory Uses:- A',
    'C': 'Conditional Uses:- C',
    'SCU': 'See Code Uses:- SCU',
    'T': 'Temporary Uses:- T',
    'AD': 'Administrative Uses:- AD',
    'L': 'Limited Uses:- L'
}

def validate_city_folder_name(city_folder: str, shp_name: str) -> Tuple[bool, str]:
    """Validate city folder name matches SHP file naming convention"""
    try:
        parts = shp_name.lower().split('_')
        if len(parts) < 2:
            return False, "Invalid SHP file name format"
        
        state_abbr = parts[0].upper()
        city_parts = parts[1:]
        expected_city = ''.join([part.capitalize() for part in city_parts]) + state_abbr
        
        return city_folder == expected_city, f"Expected folder name: {expected_city}"
    except Exception as e:
        return False, f"Validation error: {str(e)}"

def validate_file_naming(city_folder_path: str, shp_name: str) -> Dict[str, List[str]]:
    """Validate all file names match the SHP naming convention"""
    errors = {}
    base_name = shp_name.lower()
    
    # Check PLU files
    plu_folder = os.path.join(city_folder_path, "PLU")
    if os.path.exists(plu_folder):
        plu_files = [f for f in os.listdir(plu_folder) if f.endswith(".csv")]
        for file in plu_files:
            if not file.lower().startswith(base_name + "_plu_"):
                errors.setdefault("PLU", []).append(file)
    
    # Check controls file
    controls_files = [f for f in os.listdir(city_folder_path) 
                     if "_controls.csv" in f.lower() and not f.lower().startswith(base_name)]
    if controls_files:
        errors["Controls"] = controls_files
    
    # Check zone type files
    zone_type_files = [f for f in os.listdir(city_folder_path) 
                      if "_zone_type" in f.lower() and not f.lower().startswith(base_name)]
    if zone_type_files:
        errors["Zone Type"] = zone_type_files
    
    return errors


def get_shp_file(city_folder_path: str) -> Tuple[Optional[str], Optional[str]]:
    """Find and return the main SHP file path"""
    # Look for SHP folder in the temp directory
    shp_folder = os.path.join(city_folder_path, "SHP")
    if not os.path.exists(shp_folder):
        # Try alternate path structure (if zip contained folder)
        for root, dirs, files in os.walk(city_folder_path):
            if "SHP" in dirs:
                shp_folder = os.path.join(root, "SHP")
                break

    if not os.path.exists(shp_folder):
        return None, "SHP folder not found"

    shp_files = [f for f in os.listdir(shp_folder) if f.lower().endswith('.shp')]
    if not shp_files:
        return None, "No SHP files found in SHP folder"

    main_shp = shp_files[0]
    return os.path.join(shp_folder, main_shp), None

def get_city_folder(plu_folder_path: str) -> str:
    """Navigate up from PLU folder to find the city folder"""
    if plu_folder_path.endswith("PLU"):
        return os.path.dirname(plu_folder_path)
    return plu_folder_path

def get_shp_zone_codes(shp_path: str) -> Tuple[Set[str], Optional[str]]:
    """Get unique zone codes from SHP file"""
    try:
        gdf = gpd.read_file(shp_path)
        if 'zone_code' not in gdf.columns:
            return set(), "No 'zone_code' column found in SHP file"
        return set(gdf['zone_code'].unique()), None
    except Exception as e:
        return set(), f"Error reading SHP file: {str(e)}"

def get_csv_name(shp_path: str) -> str:
    """Generate Key_PLUS filename from SHP name"""
    shp_name = os.path.splitext(os.path.basename(shp_path))[0].lower()
    return f"{shp_name}_Key_PLUS.csv"

def create_keys_csv(folder_path: str, shp_path: str) -> Tuple[Optional[str], Optional[str]]:
    """Create empty Key_PLUS file with correct headers"""
    try:
        file_name = get_csv_name(shp_path)
        path = os.path.join(folder_path, file_name)
        with open(path, mode="w", newline="") as file:
            writer = csv.writer(file)
            writer.writerow(["zone_code", "missing_zone"])
        return path, None
    except Exception as e:
        return None, f"Error creating Key_PLUS file: {str(e)}"

def write_in_keys_csv(folder_path: str, shp_path: str, unique_zone_codes: Set[str], 
                     plu_folder_path: str, missing_zones: Set[str]) -> Tuple[Optional[str], Optional[str]]:
    """Write data to Key_PLUS file with proper formatting"""
    try:
        file_name = get_csv_name(shp_path)
        csv_filename = os.path.join(folder_path, file_name)
        
        # Get all unique PLU use types
        unique_use_types = get_all_plu_use_types(plu_folder_path)
        
        with open(csv_filename, 'w', newline='') as csv_file:
            writer = csv.writer(csv_file)
            writer.writerow(["zone_code", "missing_zone"])
            
            for item in sorted(unique_zone_codes):
                writer.writerow([item, 'missing' if item in missing_zones else ''])
            
            writer.writerow([""])  # Empty line separator
            writer.writerow([""])  # Empty line separator
            
            # Write sorted unique use types
            for use_type in sorted(unique_use_types):
                writer.writerow([use_type])
                
        return csv_filename, None
    except Exception as e:
        return None, f"Error writing to Key_PLUS file: {str(e)}"
    
def get_all_plu_use_types(plu_folder_path: str) -> Set[str]:
    """Get all unique PLU use types from all PLU files in the folder"""
    unique_use_types = set()
    
    for file_name in os.listdir(plu_folder_path):
        if file_name.endswith(".csv") and "_PLU_" in file_name:
            file_path = os.path.join(plu_folder_path, file_name)
            try:
                with open(file_path, 'r', encoding='utf-8') as f:
                    reader = csv.reader(f)
                    next(reader)  # Skip header
                    for row in reader:
                        if len(row) > 1:  # Check if row has at least 2 columns
                            for cell in row[1:]:  # Skip first column (use types)
                                if cell.strip() and cell.strip() in PLU_USE_TYPES:
                                    unique_use_types.add(PLU_USE_TYPES[cell.strip()])
            except Exception as e:
                st.error(f"Error reading {file_name}: {str(e)}")
    
    return unique_use_types

def clean_plu_text(text: str, is_first_column: bool = False) -> Tuple[str, List[Tuple[str, str]]]:
    """Enhanced cleaning with all new features"""
    spelling_errors = []
    
    if not isinstance(text, str):
        return text, spelling_errors
    
    # 1. Remove commas between numbers (20,000 → 20000)
    text = re.sub(r'(?<=\d),(?=\d)', '', text)
    
    # 2. Convert "Sq ft" to "Square feet"
    text = re.sub(r'\bSq\s*ft\b', 'Square feet', text, flags=re.IGNORECASE)
    
    # 4. Remove comma at end of sentence
    text = re.sub(r',\s*$', '', text)
    
    if is_first_column:
        # Original cleaning rules
        text = re.sub(r'[;:"\.]', '', text)
        replacements = [
            (r'\s*/\s*', ' or '),
            (r'\s*&\s*', ' and '),
            (r'\s*-\s*', ' '),
            (r'<', 'less than '),
            (r'>', 'greater than '),
            (r'\s+', ' '),
            (r'\s*%\s*', ' percent '),
        ]
        for pattern, repl in replacements:
            text = re.sub(pattern, repl, text)
        
        text = text.strip()
        if text:
            text = text[0].upper() + text[1:].lower()
            
            # 3. Spell Check (first column only)
            words = re.findall(r'\b\w+\b', text)
            misspelled = spell.unknown(words)
            for word in misspelled:
                correct_word = spell.correction(word)
                if correct_word and correct_word.lower() != word.lower():
                    spelling_errors.append((word, correct_word))
    
    return text, spelling_errors

def correct_plu_files(plu_folder_path: str) -> Dict[str, Dict[str, List]]:
    """Enhanced correction with all new features"""
    results = {}
    plu_files = [f for f in os.listdir(plu_folder_path) if f.endswith(".csv") and "_PLU_" in f]

    for file_name in plu_files:
        file_path = os.path.join(plu_folder_path, file_name)
        temp_path = os.path.join(plu_folder_path, f"temp_{file_name}")
        changes = []
        spelling_issues = []

        try:
            with open(file_path, 'r', encoding='utf-8') as infile, \
                 open(temp_path, 'w', newline='', encoding='utf-8') as outfile:

                reader = csv.reader(infile)
                writer = csv.writer(outfile)
                header = next(reader)
                writer.writerow(header)

                for row_num, row in enumerate(reader, start=2):
                    original_row = row.copy()
                    cleaned_row = []
                    row_spelling_issues = []

                    for col_num, cell in enumerate(row):
                        cleaned_cell, cell_errors = clean_plu_text(cell, is_first_column=(col_num == 0))
                        cleaned_row.append(cleaned_cell)
                        if cell_errors:
                            row_spelling_issues.extend([(row_num, word, correction) for word, correction in cell_errors])

                    if cleaned_row != original_row:
                        changes.append(f"Row {row_num}: {original_row} → {cleaned_row}")
                    if row_spelling_issues:
                        spelling_issues.extend(row_spelling_issues)

                    writer.writerow(cleaned_row)

            os.replace(temp_path, file_path)
            results[file_name] = {
                'changes': changes,
                'spelling_issues': spelling_issues
            }

        except Exception as e:
            if os.path.exists(temp_path):
                os.remove(temp_path)
            results[file_name] = {'error': str(e)}

    st.session_state.plu_files_corrected = True
    return results

def similar(a: str, b: str) -> float:
    """Calculate similarity ratio between two strings"""
    return SequenceMatcher(None, a, b).ratio()

def find_similar_duplicates(uses: List[str], threshold: float = 0.98) -> List[Dict[str, Union[str, List[int]]]]:
    """Find similar uses that are nearly duplicates"""
    duplicates = []
    seen = {}
    
    for idx, use in enumerate(uses):
        use = use.strip()
        if not use:
            continue
            
        found = False
        for seen_use, seen_indices in seen.items():
            if similar(use, seen_use) >= threshold:
                seen[seen_use].append(idx + 1)  # +1 for 1-based row numbering
                found = True
                break
                
        if not found:
            seen[use] = [idx + 1]
    
    for use, indices in seen.items():
        if len(indices) > 1:
            duplicates.append({'use': use, 'rows': indices})
    
    return duplicates

def check_duplicate_uses_in_column(csv_path: str, column_index: int = 0) -> Union[List[Dict[str, Union[str, List[int]]]], Dict[str, str]]:
    """Check for duplicate and similar uses in the specified column of a CSV file"""
    try:
        with open(csv_path, 'r', encoding='utf-8') as file:
            reader = csv.reader(file)
            next(reader)  # Skip header
            rows = list(reader)
            
            # Extract uses from the specified column
            uses = [row[column_index] for row in rows if len(row) > column_index]
            
            # Find exact duplicates
            exact_duplicates = []
            seen = {}
            for row_num, row in enumerate(rows, start=2):
                if len(row) > column_index:
                    value = row[column_index].strip()
                    if value:
                        if value in seen:
                            seen[value].append(row_num)
                        else:
                            seen[value] = [row_num]
            
            for value, row_nums in seen.items():
                if len(row_nums) > 1:
                    exact_duplicates.append({'use': value, 'rows': row_nums})
            
            # Find similar duplicates (98% threshold)
            similar_duplicates = find_similar_duplicates(uses)
            
            # Combine results
            return {
                'exact_duplicates': exact_duplicates,
                'similar_duplicates': similar_duplicates
            }
            
    except Exception as e:
        return {'error': f"Error processing file: {str(e)}"}

def validate_plu_use_types(plu_folder_path: str) -> Dict[str, List[Tuple[int, int, str]]]:
    """Validate PLU use types (including trailing spaces)"""
    invalid_uses = {}
    valid_types = {'S', 'P', 'A', 'C', 'SCU', 'T', 'AD', 'L'}
    
    for file_name in [f for f in os.listdir(plu_folder_path) if f.endswith(".csv") and "_PLU_" in f]:
        file_path = os.path.join(plu_folder_path, file_name)
        errors = []
        
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                reader = csv.reader(f)
                next(reader)  # Skip header
                
                for row_num, row in enumerate(reader, start=2):
                    for col_num, value in enumerate(row[1:], start=2):  # Skip first column
                        val = value.strip()  # First strip to check content
                        original_val = value.rstrip()  # Keep original for space check
                        
                        # Check for trailing spaces
                        if len(original_val) != len(value):
                            errors.append((row_num, col_num, f"'{value}' (has trailing spaces)"))
                        # Check for invalid type (after verifying it's not just whitespace)
                        elif val and val not in valid_types:
                            errors.append((row_num, col_num, val))
        
        except Exception as e:
            errors.append((0, 0, f"File Error: {str(e)}"))
        
        if errors:
            invalid_uses[file_name] = errors
    
    return invalid_uses

def validate_controls_file(controls_path: str, shp_zone_codes: Set[str]) -> List[str]:
    """Validate controls file structure and content"""
    errors = []
    try:
        with open(controls_path, 'r') as file:
            reader = csv.reader(file)
            header = next(reader)
            
            # Check for duplicate columns
            seen_cols = {}
            for idx, col in enumerate(header):
                clean_col = col.replace('ï»¿', '')  # Handle UTF-8 BOM
                if clean_col in seen_cols:
                    errors.append(f"Duplicate column '{clean_col}' at positions {seen_cols[clean_col]+1} and {idx+1}")
                else:
                    seen_cols[clean_col] = idx
            
            # Check zone codes match SHP
            zone_code_col = None
            for i, col in enumerate(header):
                if col.replace('ï»¿', '') == 'zone_code':
                    zone_code_col = i
                    break
            
            if zone_code_col is None:
                errors.append("Missing 'zone_code' column in controls file")
            else:
                controls_zone_codes = set()
                for row in reader:
                    if len(row) > zone_code_col:
                        controls_zone_codes.add(row[zone_code_col])
                
                extra_zones = controls_zone_codes - shp_zone_codes
                if extra_zones:
                    errors.append(f"Extra zones in controls: {', '.join(sorted(extra_zones))}")
                
                missing_zones = shp_zone_codes - controls_zone_codes
                if missing_zones:
                    errors.append(f"Missing zones in controls: {', '.join(sorted(missing_zones))}")
                
            # Check standard columns
            file.seek(0)
            reader = csv.reader(file)
            header = [col.replace('ï»¿', '') for col in next(reader)]
            
            if header[:16] != STANDARD_CONTROLS_COLUMNS:
                errors.append("First 16 columns don't match standard controls format")
            
            # Check values
            for row_num, row in enumerate(reader, 1):
                for col_num, value in enumerate(row[1:16], 1):
                    if value and not (value.isdigit() or is_valid_decimal(value) or value in ["NA", "STF"]):
                        errors.append(f"Invalid value at row {row_num}, column {col_num+1}: {value}")
            
            # Check blank cells
            file.seek(0)
            reader = csv.reader(file)
            blank_cells = []
            for row_idx, row in enumerate(reader):
                for col_idx, cell in enumerate(row):
                    if cell == '':
                        blank_cells.append((row_idx+1, col_idx+1))
            if blank_cells:
                errors.append(f"Blank cells found at: {blank_cells}")
                
    except Exception as e:
        errors.append(f"Error processing controls file: {str(e)}")
    
    return errors

def is_valid_decimal(input_string: str) -> bool:
    """Check if string is a valid decimal number"""
    if not input_string or input_string == '.':
        return False
    if input_string[0] in '+-':
        input_string = input_string[1:]
    return input_string.replace('.', '', 1).isdigit()


def validate_zone_type_file(file_path: str) -> List[str]:
    """Validate zone type file structure and content"""
    errors = []
    try:
        df = pd.read_csv(file_path)
        
        # Determine required columns
        if 'zone_code_link' in df.columns and df['zone_code_link'].str.contains('zoneomics').any():
            required_columns = ZONEOMICS_REQUIRED_COLUMNS
        else:
            required_columns = FULL_REQUIRED_COLUMNS
        
        # Check column sequence
        actual_columns = df.columns.tolist()
        zone_guide_present = 'zone_guide' in actual_columns
        filtered_required = [col for col in required_columns if col != 'zone_guide' or zone_guide_present]
        
        if actual_columns[:len(filtered_required)] != filtered_required:
            errors.append(f"Column sequence error. Expected: {filtered_required}")
        
        # Validate zone types
        if 'zone_type' in df.columns and 'zone_sub_type' in df.columns:
            for idx, row in df.iterrows():
                if row['zone_type'] not in VALID_ZONES:
                    errors.append(f"Row {idx+1}: Invalid zone_type '{row['zone_type']}'")
                elif row['zone_sub_type'] not in VALID_ZONES[row['zone_type']]:
                    errors.append(f"Row {idx+1}: Invalid zone_sub_type '{row['zone_sub_type']}' for zone_type '{row['zone_type']}'")
        else:
            errors.append("Missing zone_type or zone_sub_type columns")
            
    except Exception as e:
        errors.append(f"Error processing file: {str(e)}")
    
    return errors

def perform_plu_controls_qa(plu_folder_path: str) -> Dict:
    """Perform comprehensive QA on PLU and controls files"""
    results = {
        'shp_zone_codes': set(),
        'plu_zone_codes': set(),
        'missing_zones': set(),
        'plu_issues': [],
        'duplicate_uses': {},
        'similar_duplicates': {},
        'controls_issues': [],
        'zone_type_issues': {},
        'status_messages': [],
        'key_file_created': False,
        'key_file_path': None,
        'all_unique_records': set(),
        'folder_valid': False,
        'naming_errors': {}
    }
    
    try:
        city_folder = get_city_folder(plu_folder_path)
        city_folder_name = os.path.basename(city_folder)
        
        # Validate city folder and SHP file
        shp_path, shp_error = get_shp_file(city_folder)
        if not shp_path:
            results['status_messages'].append(("SHP Validation", f"❌ {shp_error}"))
            return results
        
        shp_name = os.path.splitext(os.path.basename(shp_path))[0]
        is_valid, validation_msg = validate_city_folder_name(city_folder_name, shp_name)
        
        if not is_valid:
            results['status_messages'].append(("Folder Validation", f"❌ {validation_msg}"))
            return results
        
        results['folder_valid'] = True
        results['status_messages'].append(("Folder Validation", "✅ Valid folder structure"))

        naming_errors = validate_file_naming(city_folder, shp_name)
        if naming_errors:
            results['naming_errors'] = naming_errors
            for file_type, files in naming_errors.items():
                results['status_messages'].append(
                    (f"{file_type} Naming", 
                    f"❌ {len(files)} files don't match SHP naming: {', '.join(files)}"))
        else:
            results['status_messages'].append(
                ("File Naming", "✅ All files follow SHP naming convention"))
        
        
        # Get zone codes from SHP
        zone_codes, zone_error = get_shp_zone_codes(shp_path)
        if zone_error:
            results['status_messages'].append(("SHP Zone Codes", f"❌ {zone_error}"))
        else:
            results['shp_zone_codes'] = zone_codes
            results['status_messages'].append(("SHP Zone Codes", f"✅ Found {len(zone_codes)} zone codes"))
        
        # Process PLU files
        try:
            plu_files = [f for f in os.listdir(plu_folder_path) if f.endswith(".csv") and "_PLU_" in f]
            if not plu_files:
                results['status_messages'].append(("PLU Files", "⚠️ No PLU files found"))
            else:
                results['status_messages'].append(("PLU Files", f"✅ Found {len(plu_files)} PLU files"))
                
                for file_name in plu_files:
                    file_path = os.path.join(plu_folder_path, file_name)
                    duplicates_result = check_duplicate_uses_in_column(file_path)
                    
                    if isinstance(duplicates_result, dict) and 'error' in duplicates_result:
                        results['plu_issues'].append({'file': file_name, 'issues': [duplicates_result['error']]})
                        results['status_messages'].append((f"PLU {file_name}", f"❌ {duplicates_result['error']}"))
                    else:
                        if duplicates_result['exact_duplicates']:
                            results['duplicate_uses'][file_name] = duplicates_result['exact_duplicates']
                        if duplicates_result['similar_duplicates']:
                            results['similar_duplicates'][file_name] = duplicates_result['similar_duplicates']
                        
                        status_msg = []
                        if duplicates_result['exact_duplicates']:
                            status_msg.append(f"{len(duplicates_result['exact_duplicates'])} exact duplicates")
                        if duplicates_result['similar_duplicates']:
                            status_msg.append(f"{len(duplicates_result['similar_duplicates'])} similar duplicates")
                        
                        if status_msg:
                            results['status_messages'].append((f"PLU {file_name}", f"⚠️ Found {' and '.join(status_msg)}"))
                        else:
                            results['status_messages'].append((f"PLU {file_name}", "✅ No duplicates found"))
                    
                    # Collect zone codes from PLU files
                    try:
                        with open(file_path, 'r', encoding='utf-8') as csvfile:
                            reader = csv.reader(csvfile)
                            header = next(reader)
                            if len(header) > 1:  # Skip if only one column
                                results['plu_zone_codes'].update(set(header[1:]))  # Skip first column (use types)
                    except Exception as e:
                        results['plu_issues'].append({'file': file_name, 'issues': [f"Error reading file: {str(e)}"]})
        
        except Exception as e:
            results['status_messages'].append(("PLU Files Processing", f"❌ Error processing PLU files: {str(e)}"))
        
        # Get all unique PLU use types
        results['all_unique_records'] = validate_plu_use_types(plu_folder_path)
        
        # Create Key_PLUS file if we have SHP zone codes
        if results['shp_zone_codes']:
            results['missing_zones'] = results['shp_zone_codes'] - results['plu_zone_codes']
            
            key_path, key_error = create_keys_csv(city_folder, shp_path)
            if key_error:
                results['status_messages'].append(("Key_PLUS File", f"❌ {key_error}"))
            else:
                final_path, write_error = write_in_keys_csv(
                    city_folder, shp_path, 
                    results['shp_zone_codes'], 
                    plu_folder_path, 
                    results['missing_zones']
                )
                
                if write_error:
                    results['status_messages'].append(("Key_PLUS File", f"❌ {write_error}"))
                else:
                    results['key_file_path'] = final_path
                    results['key_file_created'] = True
                    results['status_messages'].append(("Key_PLUS File", f"✅ Created at {final_path}"))
        
        # Add this right after Key_PLUS file creation
        try:
            invalid_uses = validate_plu_use_types(plu_folder_path)
            results['invalid_plu_uses'] = invalid_uses
    
            if invalid_uses:
                error_count = sum(len(errors) for errors in invalid_uses.values())
                results['status_messages'].append(
                    ("PLU Use Types", f"❌ Found {error_count} invalid use types")
                )
            else:
                results['status_messages'].append(
                    ("PLU Use Types", "✅ No issues in PLU Use Types")
                )
        
        except Exception as e:
            results['status_messages'].append(
                ("PLU Use Types", f"❌ Validation failed: {str(e)}")
            )
            results['invalid_plu_uses'] = {}

        
        # Check controls file
        controls_file = next((f for f in os.listdir(city_folder) if '_controls.csv' in f.lower()), None)
        if controls_file:
            controls_path = os.path.join(city_folder, controls_file)
            results['controls_issues'] = validate_controls_file(controls_path, results['shp_zone_codes'])
            
            if results['controls_issues']:
                results['status_messages'].append(("Controls File", f"⚠️ Found {len(results['controls_issues'])} issues"))
            else:
                results['status_messages'].append(("Controls File", "✅ No issues found"))
        else:
            results['status_messages'].append(("Controls File", "⚠️ No controls file found"))
        
        # Check zone type files
        zone_type_files = [f for f in os.listdir(city_folder) if 'zone_type' in f.lower() and f.endswith('.csv')]
        if not zone_type_files:
            results['status_messages'].append(("Zone Type Files", "⚠️ No zone type files found"))
        else:
            results['status_messages'].append(("Zone Type Files", f"✅ Found {len(zone_type_files)} files"))
            
            for file_name in zone_type_files:
                file_path = os.path.join(city_folder, file_name)
                issues = validate_zone_type_file(file_path)
                
                if issues:
                    results['zone_type_issues'][file_name] = issues
                    results['status_messages'].append((f"Zone Type {file_name}", f"⚠️ Found {len(issues)} issues"))
                else:
                    results['status_messages'].append((f"Zone Type {file_name}", "✅ No issues found"))
    
    except Exception as e:
        results['status_messages'].append(("QA Process", f"❌ Critical error: {str(e)}"))
    
    return results

def create_word_report(qa_results: Dict, folder_path: str) -> str:
    """Create a Word document report of QA results"""
    doc = Document()
    
    # Title
    title = doc.add_heading('Zoning Data QA Report', 0)
    title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    
    # Report Metadata
    doc.add_paragraph(f"Date: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    doc.add_paragraph(f"Folder Path: {folder_path}")
    doc.add_paragraph()
    
    # Status Summary
    doc.add_heading('QA Status Summary', level=1)
    for step, status in qa_results.get('status_messages', []):
        p = doc.add_paragraph()
        p.add_run(f"{step}: ").bold = True
        p.add_run(status)

    # File Naming Validation Section
    doc.add_heading('File Naming Validation', level=1)
    if qa_results.get('naming_errors'):
        doc.add_paragraph('Files with incorrect naming convention:', style='Heading 2')
        for file_type, files in qa_results['naming_errors'].items():
            doc.add_paragraph(f"{file_type} Files:", style='Heading 3')
            for file in files:
                doc.add_paragraph(f"• {file}", style='List Bullet')
    else:
        doc.add_paragraph('All files follow correct naming convention', style='Intense Quote')
    
    
    # PLU/Controls Section
    doc.add_heading('PLU/Controls QA Results', level=1)
    # Add in the PLU/Controls section of create_word_report()
    doc.add_heading('PLU Use Types Validation', level=2)
    if qa_results.get('invalid_plu_uses'):
        doc.add_paragraph('Invalid PLU Use Types Found:', style='Heading 3')
        for file_name, errors in qa_results['invalid_plu_uses'].items():
            doc.add_paragraph(f"File: {file_name}", style='Heading 4')
            table = doc.add_table(rows=1, cols=3)
            table.style = 'LightShading-Accent1'
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = 'Row'
            hdr_cells[1].text = 'Column'
            hdr_cells[2].text = 'Invalid Type'
        
            for row, col, err_type in errors:
                row_cells = table.add_row().cells
                row_cells[0].text = str(row)
                row_cells[1].text = str(col)
                row_cells[2].text = err_type
    else:
        doc.add_paragraph('All PLU use types are valid', style='Intense Quote')
    
    # SHP Zone Codes
    doc.add_heading('Zone Codes from SHP File', level=2)
    if qa_results.get('shp_zone_codes'):
        doc.add_paragraph(', '.join(sorted(qa_results['shp_zone_codes'])))
    else:
        doc.add_paragraph('No SHP zone codes found', style='Intense Quote')
    
    # Missing Zones
    doc.add_heading('Missing Zones (In SHP but not in PLU files)', level=2)
    if qa_results.get('missing_zones'):
        for zone in sorted(qa_results['missing_zones']):
            doc.add_paragraph(zone)
    else:
        doc.add_paragraph('No missing zones found', style='Intense Quote')
    
    # PLU Key File
    doc.add_heading('PLU Key File', level=2)
    if qa_results.get('key_file_created'):
        doc.add_paragraph(f"Created at: {qa_results['key_file_path']}")
    else:
        doc.add_paragraph('PLU key file was not created', style='Intense Quote')
    
    # PLU File Issues
    doc.add_heading('PLU File Issues', level=2)
    if qa_results.get('plu_issues'):
        for file_issues in qa_results['plu_issues']:
            doc.add_heading(file_issues['file'], level=3)
            for issue in file_issues['issues']:
                doc.add_paragraph(f"• {issue}")
    else:
        doc.add_paragraph('No general PLU file issues found', style='Intense Quote')
    
    # PLU Duplicates
    if qa_results.get('duplicate_uses'):
        doc.add_heading('EXACT DUPLICATE USES FOUND IN PLU FILES', level=2)
        for file_name, duplicates in qa_results['duplicate_uses'].items():
            doc.add_heading(f"In file: {file_name}", level=3)
            for dup in duplicates:
                doc.add_paragraph(f"• Use '{dup['use']}' appears in rows: {', '.join(map(str, dup['rows']))}")
    else:
        doc.add_paragraph('No exact duplicate uses found in PLU files', style='Intense Quote')
    
    # PLU Similar Duplicates
    if qa_results.get('similar_duplicates'):
        doc.add_heading('SIMILAR DUPLICATE USES FOUND IN PLU FILES (98% match)', level=2)
        for file_name, duplicates in qa_results['similar_duplicates'].items():
            doc.add_heading(f"In file: {file_name}", level=3)
            for dup in duplicates:
                doc.add_paragraph(f"• Use '{dup['use']}' appears in rows: {', '.join(map(str, dup['rows']))}")
    else:
        doc.add_paragraph('No similar duplicate uses found in PLU files', style='Intense Quote')
    
    # Controls Issues
    doc.add_heading('Controls File Issues', level=2)
    if qa_results.get('controls_issues'):
        for issue in qa_results['controls_issues']:
            doc.add_paragraph(f"• {issue}")
    else:
        doc.add_paragraph('No issues found in Controls file', style='Intense Quote')
    
    # Zone Type Issues
    doc.add_heading('Zone Type QA Results', level=1)
    if qa_results.get('zone_type_issues'):
        for file_name, issues in qa_results['zone_type_issues'].items():
            doc.add_heading(file_name, level=2)
            for issue in issues:
                doc.add_paragraph(f"• {issue}")
    else:
        doc.add_paragraph('No issues found in Zone Type files', style='Intense Quote')
    
    # Save to a temporary file
    report_path = os.path.join(os.getcwd(), "zoning_qa_report.docx")
    doc.save(report_path)
    
    return report_path

# Streamlit UI

def extract_uploaded_zip(uploaded_zip):
    """Extract uploaded zip to temporary directory and return path"""
    if uploaded_zip is None:
        return None

    # Create temp directory
    temp_dir = tempfile.mkdtemp()

    # Save zip file temporarily
    zip_path = os.path.join(temp_dir, "uploaded.zip")
    with open(zip_path, "wb") as f:
        f.write(uploaded_zip.getbuffer())

    # Extract zip
    with zipfile.ZipFile(zip_path, 'r') as zip_ref:
        zip_ref.extractall(temp_dir)

    # Remove the zip file
    os.remove(zip_path)

    return temp_dir


# Add this cleanup function
def cleanup_temp_dir(temp_dir):
    """Remove temporary directory and its contents"""
    if temp_dir and os.path.exists(temp_dir):
        shutil.rmtree(temp_dir)

def img_to_base64(img_path):
    return base64.b64encode(Path(img_path).read_bytes()).decode()

# 3. Your logo header (replace with your actual path)
logo_path = "zoneomics_icon.png"  # Update this path
logo_base64 = img_to_base64(logo_path)

# 4. Replace your st.title() with this:
st.markdown(
    f"""
    <div class="header-container">
        <img class="header-logo" src="data:image/png;base64,{logo_base64}">
        <h1 class="header-title">Zoning Data Quality Assurance Tool</h1>
    </div>
    """, 
    unsafe_allow_html=True
)
# Sidebar for folder selection
with st.sidebar:
    st.header("Configuration")
    uploaded_zip = st.file_uploader(
        "Upload Zoning Data Folder (as ZIP)",
        type="zip",
        help="Upload a zipped folder containing PLU, SHP, and other required files"
    )

    if st.button("Run QA Analysis"):
        if uploaded_zip:
            # Extract to temp folder
            temp_dir = extract_uploaded_zip(uploaded_zip)

            # Find PLU folder in extracted contents
            plu_folder_path = None
            for root, dirs, files in os.walk(temp_dir):
                if "PLU" in dirs:
                    plu_folder_path = os.path.join(root, "PLU")
                    break

            if plu_folder_path:
                st.session_state.current_folder = temp_dir
                with st.spinner("Performing QA analysis..."):
                    st.session_state.qa_results = perform_plu_controls_qa(plu_folder_path)
                    st.session_state.temp_dir = temp_dir  # Store for cleanup later
                st.success("QA analysis completed!")
            else:
                st.error("No PLU folder found in uploaded zip")
                cleanup_temp_dir(temp_dir)
        else:
            st.error("Please upload a zip file first")

# Main content area - Initial instructions
if not st.session_state.current_folder:
    st.markdown("""
    ## Welcome to the Zoning Data QA Tool
    
    This tool performs comprehensive quality assurance checks on:
    - PLU (Permitted Land Use) files
    - Controls files (zoning standards)
    - Zone Type files
    
    ### How to use:
    1. Enter the path to your PLU folder in the sidebar
    2. Click "Run QA Analysis"
    3. Review the QA results
    4. Use the "Correct PLU Files" option if needed
    
    Checks include:
    - Duplicate uses in PLU files
    - Zone code consistency between SHP and PLU files
    - Missing zone codes
    - Controls file validation
    - Zone type validation
    - Data completeness
    - PLU Corrections
    """)
    st.markdown(
        '<div class="footer">Powered by Zoneomics ©</div>',
        unsafe_allow_html=True
    )
# Show results after analysis
if st.session_state.current_folder and st.session_state.qa_results:
    st.markdown(f"### QA Results for: `{st.session_state.current_folder}`")
    
    # Status Overview (only once)
    with st.expander("QA Status Overview", expanded=True):
        for step, status in st.session_state.qa_results.get('status_messages', []):
            if status.startswith("✅"):
                st.success(f"{step}: {status}")
            elif status.startswith("⚠️"):
                st.warning(f"{step}: {status}")
            elif status.startswith("❌"):
                st.error(f"{step}: {status}")
            else:
                st.info(f"{step}: {status}")
    # Naming Errors Section
    with st.expander("File Naming Validation", expanded=False):
        if naming_errors := st.session_state.qa_results.get('naming_errors'):
            st.error("Files with incorrect naming convention:")
        
            for file_type, files in naming_errors.items():
                st.markdown(f"**{file_type} Files:**")
                for file in files:
                    st.markdown(f"- `{file}`")
                st.markdown("---")
        
            st.warning("All files should follow the SHP file naming pattern. Example:")
            st.code("tn_alexandria_controls.csv  # Right Naming Convention\ntn_alexandri_controls.csv   # Wrong (missing 'a')")
        else:
            st.success("All files follow the correct naming convention")
    

    
    # ---- PLU Use Types Validation Section ----
    with st.expander("PLU Use Types Validation", expanded=False):
        if 'qa_results' not in st.session_state:
            st.info("Run QA analysis to validate PLU use types")
        else:
            if invalid_results := st.session_state.qa_results.get('invalid_plu_uses'):
                # Summary stats
                cols = st.columns(3)
                cols[0].metric("Files with Issues", len(invalid_results))
                cols[1].metric("Total Errors", sum(len(e) for e in invalid_results.values()))
                cols[2].metric("Valid Types", len(PLU_USE_TYPES))
            
                # File selector
                selected_file = st.selectbox(
                    "Select file to view errors:",
                    options=list(invalid_results.keys()),
                    key="plu_use_file_select"
                )
            
                # Error details
                st.markdown(f"**Errors in {selected_file}:**")
                error_df = pd.DataFrame(
                    [(row, col, err) for row, col, err in invalid_results[selected_file]],
                    columns=["Row", "Column", "Invalid Type"]
                )
                st.table(error_df)
            else:
                st.success("All PLU use types are valid")
        
            # Valid types reference (always shown)
            st.markdown("**Valid PLU Use Types:**")
            st.table(pd.DataFrame(
                [(code, desc) for code, desc in PLU_USE_TYPES.items()],
                columns=["Code", "Description"]
            ))
    
    # Correct PLU Files Section
    if st.session_state.get('plu_files_corrected'):
        st.success("PLU files were already corrected!")
    else:
        correction_expander = st.expander("Correct PLU Files", expanded=False)
        with correction_expander:
            st.markdown("""
        **Features to be corrected:**
        - Converting to proper sentence case (only first column)
        - Removing special characters (; : .)
        - Replacing "/" with "or"
        - Replacing "&" with "and"
        - Replacing "-" with space
        - Replacing "<" and ">" with words
        - Removing double spaces
        - Replacing "%" with "percent 
        - Remove commas in numbers (20,000 → 20000)
        - Convert 'Sq ft' to 'Square feet'
        - Spell check with suggestions
        - Remove trailing commas
        """)

            if st.button("Run PLU Corrections"):
                with st.spinner("Applying enhancements..."):
                    correction_results = correct_plu_files(st.session_state.current_folder)

                st.session_state.correction_results = correction_results
                st.session_state.plu_files_corrected = True
                st.success("Corrections applied successfully!")

            # Display results in a clean way without nesting
                st.subheader("Correction Results")
            
            # Use tabs for each file
                file_tabs = st.tabs([f"📄 {name}" for name in correction_results.keys()])
            
                for tab, (file_name, result) in zip(file_tabs, correction_results.items()):
                    with tab:
                        if 'error' in result:
                            st.error(f"Error: {result['error']}")
                            continue
                    
                    # Formatting changes
                        if result['changes']:
                            st.markdown("**Formatting Changes:**")
                            st.table(pd.DataFrame(result['changes'], columns=["Changes"]))
                    
                    # Spelling suggestions
                        if result['spelling_issues']:
                            st.markdown("**Spelling Suggestions:**")
                            df = pd.DataFrame(
                                [(row, word, suggestion) for row, word, suggestion in result['spelling_issues']],
                                columns=["Row", "Word", "Suggestion"]
                            )
                            st.table(df)
                    
                        if not result['changes'] and not result['spelling_issues']:
                            st.info("No changes needed for this file")
    
    # SHP Zone Codes
    with st.expander("Zone Codes from SHP File", expanded=False):
        if st.session_state.qa_results.get('shp_zone_codes'):
            cols = st.columns(4)
            for idx, zone in enumerate(sorted(st.session_state.qa_results['shp_zone_codes'])):
                cols[idx % 4].info(zone)
        else:
            st.warning("No SHP zone codes found")
    
    # Missing Zones
    with st.expander("Missing Zones (Present in SHP but not in PLU files)", expanded=False):
        if st.session_state.qa_results.get('missing_zones'):
            st.error(f"Found {len(st.session_state.qa_results['missing_zones'])} missing zone(s):")
            cols = st.columns(4)
            for idx, zone in enumerate(sorted(st.session_state.qa_results['missing_zones'])):
                cols[idx % 4].warning(zone)
        else:
            st.success("No missing zones found - all SHP zones are present in PLU files")
    
    # Duplicate Uses Section (with tabs)
    with st.expander("Duplicate Uses in PLU Files", expanded=False):
        if st.session_state.qa_results.get('duplicate_uses') or st.session_state.qa_results.get('similar_duplicates'):
            tab1, tab2 = st.tabs(["Exact Duplicates", "Similar Duplicates"])
        
        # Exact Duplicates Tab
            with tab1:
                if st.session_state.qa_results.get('duplicate_uses'):
                    st.error("Exact duplicates found (identical text):")
                    for file_name, duplicates in st.session_state.qa_results['duplicate_uses'].items():
                        st.markdown(f"**File:** `{file_name}`")
                        for dup in duplicates:
                            st.markdown(f"- **Use:** `{dup['use']}`")
                            st.markdown(f"  - **Appears in rows:** {', '.join(map(str, dup['rows']))}")
                        st.markdown("---")
                else:
                    st.success("No exact duplicates found")
        
        # Similar Duplicates Tab
            with tab2:
                if st.session_state.qa_results.get('similar_duplicates'):
                    st.warning("Similar duplicates found (>98% match):")
                    for file_name, duplicates in st.session_state.qa_results['similar_duplicates'].items():
                        st.markdown(f"**File:** `{file_name}`")
                        for dup in duplicates:
                            st.markdown(f"- **Use:** `{dup['use']}`")
                            st.markdown(f"  - **Appears in rows:** {', '.join(map(str, dup['rows']))}")
                        st.markdown("---")
                else:
                    st.success("No similar duplicates found")
        else:
            st.success("No duplicate uses found in PLU files")
    
    # Controls Issues
    with st.expander("Controls File Issues", expanded=False):
        if st.session_state.qa_results.get('controls_issues'):
            for issue in st.session_state.qa_results['controls_issues']:
                st.error(issue)
        else:
            st.success("No issues found in Controls file")
    
    # Zone Type Issues
    with st.expander("Zone Type QA Results", expanded=False):
        if st.session_state.qa_results.get('zone_type_issues'):
            for file_name, issues in st.session_state.qa_results['zone_type_issues'].items():
                st.error(f"File: {file_name}")
                for issue in issues:
                    st.write(f"- {issue}")
        else:
            st.success("No issues found in Zone Type files")
    
    # Generate Report
    st.markdown('<div class="header-style">Generate QA Report</div>', unsafe_allow_html=True)
    
    if st.button("Generate Word Report"):
        report_path = create_word_report(st.session_state.qa_results, st.session_state.current_folder)
        
        with open(report_path, "rb") as file:
            st.download_button(
                label="Download Word Report",
                data=file,
                file_name="zoning_qa_report.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )
    st.markdown(
    '<div class="footer">Powered by Zoneomics ©</div>',
    unsafe_allow_html=True
)

